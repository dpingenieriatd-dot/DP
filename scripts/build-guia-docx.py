# Genera 19_Guia_de_Usuario.docx a partir del PDF ya construido por
# build-guia-pdf.tsx, para tener una version de Word editable siempre en
# sincronia con el PDF (que es la fuente de verdad del contenido).
#
# Uso:  npx tsx scripts/build-guia-pdf.tsx  &&  python scripts/build-guia-docx.py
#
# Reconstruye el texto por bloques y mapea el tamano de fuente a estilos de
# Word (titulos vs. cuerpo). NO reproduce el diseno grafico del PDF -para eso
# esta el PDF- pero conserva el contenido y la jerarquia, editable a mano.

import re
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor

PDF = r"C:\Users\cesar\DP\Documentacion_Proyecto\19_Guia_de_Usuario.pdf"
OUT = r"C:\Users\cesar\DP\Documentacion_Proyecto\19_Guia_de_Usuario.docx"

VERDE = RGBColor(0x27, 0x50, 0x0A)
GRIS = RGBColor(0x8A, 0x89, 0x7F)


def limpiar(t):
    # PyMuPDF puede devolver glifos de control (p. ej. la flecha que Helvetica
    # no tiene); Word rechaza cualquier caracter de control.
    return "".join(ch for ch in t if ch in "\t\n" or ch >= " ").strip()


# 1) Extraer lineas del PDF como (tipo, texto, size)
lineas = []
src = fitz.open(PDF)
for page in src:
    for b in page.get_text("dict")["blocks"]:
        if b.get("type") != 0:
            continue
        for line in b["lines"]:
            spans = [s for s in line["spans"] if s["text"].strip()]
            if not spans:
                continue
            text = limpiar("".join(s["text"] for s in spans))
            if not text:
                continue
            low = text.lower()
            if low.startswith("d&p ingenier") and "generado el" in low:
                continue  # pie de pagina
            if re.fullmatch(r"p[aá]gina \d+ de \d+", low):
                continue  # pie de pagina
            size = max(s["size"] for s in spans)
            bold = bool(spans[0]["flags"] & 0x10)
            corto = len(text) <= 62 and not text.endswith((".", ":", ",", ";"))
            es_titulo = (
                size >= 15
                or (size >= 11.5 and bold)
                or (size >= 9.8 and bold and corto and not re.match(r"\d+\.\s", text))
            )
            lineas.append(("h", text, size) if es_titulo else ("p", text, size))

# 2) Post-proceso: unir "N" + "Titulo" (numeros de seccion sueltos), quitar
#    titulos repetidos consecutivos (cabecera de pagina del PDF en cada hoja).
elems = []
i = 0
while i < len(lineas):
    kind, text, size = lineas[i]
    if kind == "h" and re.fullmatch(r"\d+", text) and i + 1 < len(lineas) and lineas[i + 1][0] == "h":
        elems.append(("h", f"{text}. {lineas[i + 1][1]}", max(size, lineas[i + 1][2])))
        i += 2
        continue
    if kind == "h" and elems and elems[-1][0] == "h" and elems[-1][1] == text:
        i += 1
        continue  # cabecera de pagina duplicada
    elems.append((kind, text, size))
    i += 1

# 3) Construir el .docx
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

visto_titulo_portada = False
for kind, text, size in elems:
    if kind == "h":
        if text == "Guia de Usuario" or text == "Guía de Usuario":
            if visto_titulo_portada:
                continue
            visto_titulo_portada = True
            p = doc.add_heading(text, level=0)
        else:
            lvl = 1 if size >= 13.5 else 2 if size >= 11.5 else 3
            p = doc.add_heading(text, level=lvl)
        for r in p.runs:
            r.font.color.rgb = VERDE
    else:
        p = doc.add_paragraph(text)
        if size <= 8.6:
            for r in p.runs:
                r.font.size = Pt(8.6)
                r.font.color.rgb = GRIS

doc.save(OUT)
print(f"DOCX guardado en {OUT}  ({len(doc.paragraphs)} parrafos)")
